#!/usr/bin/env python3
"""
generate_requirements.py

CLI tool to generate implementation requirements for developers.
Reads CSV files: Dimensions.csv, Events.csv, Implementation Packages.csv, Products.csv, SDKs.csv.

Interactive CLI: select primary business units, selection mode (Implementation Packages or Direct Events),
select packages or events, then generates requirements in Excel, PDF, and Word document.

Dependencies:
- pandas
- openpyxl
- python-docx
- reportlab
"""
import argparse
import os
import sys
import re
## Utility to clean sheet names
def sanitize_sheet_name(name):
    # Remove invalid characters for Excel sheet names and limit length
    invalid = r'[:\\/?*\[\]]'
    safe = re.sub(invalid, '_', name)
    return safe[:31]

try:
    import pandas as pd
except ImportError:
    print("Missing dependency: pandas. Install with 'pip install pandas'", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Missing dependency: python-docx. Install with 'pip install python-docx'", file=sys.stderr)
    sys.exit(1)

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
except ImportError:
    print("Missing dependency: reportlab. Install with 'pip install reportlab'", file=sys.stderr)
    sys.exit(1)

def read_csv(path):
    if not os.path.isfile(path):
        print(f"CSV file not found: {path}", file=sys.stderr)
        sys.exit(1)
    return pd.read_csv(path, encoding='utf-8-sig')

def choose_from_list(options, prompt):
    if not options:
        print(f"No options available for {prompt}", file=sys.stderr)
        sys.exit(1)
    print(f"\n{prompt}:")
    for idx, opt in enumerate(options, 1):
        print(f"{idx}. {opt}")
    print("Enter numbers separated by commas (e.g., 1,3,5), or 'a' for all:")
    choice = input("> ").strip()
    if choice.lower() in ('a', 'all'):
        return options
    selected = []
    for part in choice.split(','):
        part = part.strip()
        if not part:
            continue
        if part.isdigit():
            i = int(part)
            if 1 <= i <= len(options):
                selected.append(options[i-1])
            else:
                print(f"Ignoring invalid selection: {part}", file=sys.stderr)
        else:
            print(f"Ignoring invalid input: {part}", file=sys.stderr)
    return selected

def generate_excel(products_df, packages_df, events_df, dims_df, output_dir):
    path = os.path.join(output_dir, 'requirements.xlsx')
    with pd.ExcelWriter(path, engine='openpyxl') as writer:
        products_df.to_excel(writer, sheet_name='Products', index=False)
        if packages_df is not None:
            packages_df.to_excel(writer, sheet_name='Packages', index=False)
        events_df.to_excel(writer, sheet_name='Events', index=False)
        dims_df.to_excel(writer, sheet_name='Dimensions', index=False)
    print(f"Excel requirements written to {path}")

def generate_word(products_df, packages_df, events_df, dims_df, units, output_dir):
    path = os.path.join(output_dir, 'requirements.docx')
    doc = Document()
    doc.add_heading('Implementation Requirements', level=1)
    doc.add_heading('Business Units', level=2)
    doc.add_paragraph(', '.join(units))
    doc.add_heading('Products & SDK Settings', level=2)
    # Products table
    prod_cols = list(products_df.columns)
    table = doc.add_table(rows=1, cols=len(prod_cols))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(prod_cols):
        hdr_cells[i].text = col
    for _, row in products_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(prod_cols):
            val = row[col]
            cells[i].text = str(val) if pd.notna(val) else ''
    if packages_df is not None:
        doc.add_heading('Selected Implementation Packages', level=2)
        pkg_cols = list(packages_df.columns)
        table = doc.add_table(rows=1, cols=len(pkg_cols))
        hdr = table.rows[0].cells
        for i, col in enumerate(pkg_cols):
            hdr[i].text = col
        for _, row in packages_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(pkg_cols):
                val = row[col]
                cells[i].text = str(val) if pd.notna(val) else ''
    doc.add_heading('Events to Implement', level=2)
    evt_cols = list(events_df.columns)
    table = doc.add_table(rows=1, cols=len(evt_cols))
    hdr = table.rows[0].cells
    for i, col in enumerate(evt_cols):
        hdr[i].text = col
    for _, row in events_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(evt_cols):
            val = row[col]
            cells[i].text = str(val) if pd.notna(val) else ''
    doc.add_heading('Required Dimensions', level=2)
    dim_cols = list(dims_df.columns)
    table = doc.add_table(rows=1, cols=len(dim_cols))
    hdr = table.rows[0].cells
    for i, col in enumerate(dim_cols):
        hdr[i].text = col
    for _, row in dims_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(dim_cols):
            val = row[col]
            cells[i].text = str(val) if pd.notna(val) else ''
    doc.save(path)
    print(f"Word requirements written to {path}")

def generate_pdf(products_df, packages_df, events_df, dims_df, units, output_dir):
    path = os.path.join(output_dir, 'requirements.pdf')
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph('Implementation Requirements', styles['Heading1']))
    story.append(Spacer(1, 12))
    story.append(Paragraph('Business Units: ' + ', '.join(units), styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph('Products & SDK Settings', styles['Heading2']))
    story.append(Spacer(1, 12))
    # Products table
    prod_cols = list(products_df.columns)
    data = [prod_cols] + products_df[prod_cols].fillna('').values.tolist()
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.25, colors.black),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))
    if packages_df is not None:
        story.append(Paragraph('Selected Implementation Packages', styles['Heading2']))
        story.append(Spacer(1, 12))
        pkg_cols = list(packages_df.columns)
        data = [pkg_cols] + packages_df[pkg_cols].fillna('').values.tolist()
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('GRID', (0,0), (-1,-1), 0.25, colors.black),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 12))
    story.append(Paragraph('Events to Implement', styles['Heading2']))
    story.append(Spacer(1, 12))
    evt_cols = list(events_df.columns)
    data = [evt_cols] + events_df[evt_cols].fillna('').values.tolist()
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.25, colors.black),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))
    story.append(Paragraph('Required Dimensions', styles['Heading2']))
    story.append(Spacer(1, 12))
    dim_cols = list(dims_df.columns)
    data = [dim_cols] + dims_df[dim_cols].fillna('').values.tolist()
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.25, colors.black),
    ]))
    story.append(tbl)
    doc.build(story)
    print(f"PDF requirements written to {path}")

# --- Per-product generation helpers ---
def generate_excel_for_product(product, packages_df, events_df, dims_by_event, output_dir, event_to_pkgs):
    # Generate an Excel workbook for a single product
    import pandas as pd
    product_name = product['Name']
    safe_name = sanitize_sheet_name(product_name)
    path = os.path.join(output_dir, f'requirements_{safe_name}.xlsx')
    # Create Excel workbook per product
    with pd.ExcelWriter(path, engine='openpyxl') as writer:
        # Overview sheet: product metadata
        prod_df = product.to_frame().T.drop(columns=['Dimensions'], errors=True)
        prod_df.to_excel(writer, sheet_name='Overview', index=False)
        # If using implementation packages, create a sheet per package; else a single 'Events' sheet
        if packages_df is not None and not packages_df.empty:
            for _, pkg in packages_df.iterrows():
                pkg_name = pkg['Name']
                sheet_name = sanitize_sheet_name(pkg_name)
                # events for this package
                event_rows = []
                # find events belonging to this package
                for eid, evts in dims_by_event.items():
                    if eid in events_df['Event #'].values and pkg_name in event_to_pkgs.get(eid, []):
                        # find event metadata
                        event = events_df[events_df['Event #'] == eid].iloc[0]
                        # build rows for each dimension
                        dims = dims_by_event.get(eid)
                        if dims is not None and not dims.empty:
                            for _, drow in dims.iterrows():
                                row = {**{
                                    'Event #': event.get('Event #'),
                                    'Event Name': event.get('Event Name'),
                                    'Description': event.get('Description'),
                                    'Event Type': event.get('Event Type'),
                                    'Dev & QA Notes': event.get('Dev & QA Notes')
                                }, **drow.to_dict()}
                                event_rows.append(row)
                if not event_rows:
                    # no events for this package
                    pd.DataFrame([{'Note': 'No events selected for this package'}]).to_excel(
                        writer, sheet_name=sheet_name, index=False)
                else:
                    sheet_df = pd.DataFrame(event_rows)
                    sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
        else:
            # Direct events sheet
            rows = []
            for _, event in events_df.iterrows():
                eid = event.get('Event #')
                dims = dims_by_event.get(eid)
                if dims is not None and not dims.empty:
                    for _, drow in dims.iterrows():
                        row = {**{
                            'Event #': event.get('Event #'),
                            'Event Name': event.get('Event Name'),
                            'Description': event.get('Description'),
                            'Event Type': event.get('Event Type'),
                            'Dev & QA Notes': event.get('Dev & QA Notes')
                        }, **drow.to_dict()}
                        rows.append(row)
            if not rows:
                pd.DataFrame([{'Note': 'No direct events selected'}]).to_excel(
                    writer, sheet_name='Events', index=False)
            else:
                sheet_df = pd.DataFrame(rows)
                sheet_df.to_excel(writer, sheet_name='Events', index=False)
    print(f"Excel requirements written for {product_name} to {path}")

def generate_word_for_product(product_df, packages_df, events_df, dims_by_event, output_dir):
    # Generate a Word document for a single product
    from docx import Document
    import pandas as pd
    product_name = product_df.iloc[0]['Name']
    safe_name = sanitize_sheet_name(product_name)
    path = os.path.join(output_dir, f'requirements_{safe_name}.docx')
    doc = Document()
    doc.add_heading(f'Implementation Requirements for {product_name}', level=1)
    # Include this product's business units
    pu = product_df.iloc[0].get('primary_business_unit', '')
    su = product_df.iloc[0].get('secondary_business_unit', '')
    bu_text = []
    if pd.notna(pu) and pu:
        bu_text.append(f'Primary: {pu}')
    if pd.notna(su) and su:
        bu_text.append(f'Secondary: {su}')
    if bu_text:
        doc.add_heading('Business Units', level=2)
        doc.add_paragraph(', '.join(bu_text))
    doc.add_heading('Product & SDK Settings', level=2)
    # Product table
    prod_cols = list(product_df.columns)
    table = doc.add_table(rows=1, cols=len(prod_cols))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(prod_cols):
        hdr_cells[i].text = col
    for _, row in product_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(prod_cols):
            val = row[col]
            cells[i].text = str(val) if pd.notna(val) else ''
    # Packages
    if packages_df is not None and not packages_df.empty:
        doc.add_heading('Implementation Packages', level=2)
        for _, pkg in packages_df.iterrows():
            doc.add_heading(pkg['Name'], level=3)
            doc.add_paragraph(str(pkg.get('Description', '')))
    # Events and Dimensions
    doc.add_heading('Events & Dimensions', level=2)
    for _, event in events_df.iterrows():
        eid = event.get('Event #')
        ename = event.get('Event Name', f'Event {eid}')
        doc.add_heading(f'{ename} (Event #{eid})', level=3)
        for field in ['Description', 'Event Type', 'Priority', 'Conditions to fire', 'Dev & QA Notes']:
            if field in event and pd.notna(event[field]):
                doc.add_paragraph(f"{field}: {event[field]}")
        dims = dims_by_event.get(eid)
        if dims is not None and not dims.empty:
            doc.add_heading('Dimensions', level=4)
            dim_cols = list(dims.columns)
            tbl = doc.add_table(rows=1, cols=len(dim_cols))
            hdr = tbl.rows[0].cells
            for i, col in enumerate(dim_cols):
                hdr[i].text = col
            for _, drow in dims.iterrows():
                cells = tbl.add_row().cells
                for i, col in enumerate(dim_cols):
                    val = drow[col]
                    cells[i].text = str(val) if pd.notna(val) else ''
    doc.save(path)
    print(f"Word requirements written for {product_name} to {path}")

def generate_pdf_for_product(product_df, packages_df, events_df, dims_by_event, output_dir):
    # Generate a PDF document for a single product
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    import pandas as pd
    product_name = product_df.iloc[0]['Name']
    safe_name = sanitize_sheet_name(product_name)
    path = os.path.join(output_dir, f'requirements_{safe_name}.pdf')
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(f'Implementation Requirements for {product_name}', styles['Heading1']))
    story.append(Spacer(1, 12))
    # Include this product's business units
    pu = product_df.iloc[0].get('primary_business_unit', '')
    su = product_df.iloc[0].get('secondary_business_unit', '')
    bu_text = []
    if pd.notna(pu) and pu:
        bu_text.append(f'Primary: {pu}')
    if pd.notna(su) and su:
        bu_text.append(f'Secondary: {su}')
    if bu_text:
        story.append(Paragraph('Business Units: ' + ', '.join(bu_text), styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph('Product & SDK Settings', styles['Heading2']))
    story.append(Spacer(1, 12))
    prod_cols = list(product_df.columns)
    data = [prod_cols] + product_df[prod_cols].fillna('').values.tolist()
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.25, colors.black),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))
    if packages_df is not None and not packages_df.empty:
        story.append(Paragraph('Implementation Packages', styles['Heading2']))
        story.append(Spacer(1, 12))
        for _, pkg in packages_df.iterrows():
            story.append(Paragraph(pkg['Name'], styles['Heading3']))
            story.append(Paragraph(str(pkg.get('Description', '')), styles['Normal']))
            story.append(Spacer(1, 12))
    story.append(Paragraph('Events & Dimensions', styles['Heading2']))
    story.append(Spacer(1, 12))
    for _, event in events_df.iterrows():
        eid = event.get('Event #')
        ename = event.get('Event Name', f'Event {eid}')
        story.append(Paragraph(f'{ename} (Event #{eid})', styles['Heading3']))
        story.append(Spacer(1, 6))
        for field in ['Description', 'Event Type', 'Priority', 'Conditions to fire', 'Dev & QA Notes']:
            if field in event and pd.notna(event[field]):
                story.append(Paragraph(f"{field}: {event[field]}", styles['Normal']))
        story.append(Spacer(1, 6))
        dims = dims_by_event.get(eid)
        if dims is not None and not dims.empty:
            dim_cols = list(dims.columns)
            data = [dim_cols] + dims[dim_cols].fillna('').values.tolist()
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('GRID', (0,0), (-1,-1), 0.25, colors.black),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 12))
    doc.build(story)
    print(f"PDF requirements written for {product_name} to {path}")

def main():
    parser = argparse.ArgumentParser(description='Generate implementation requirements from CSVs.')
    parser.add_argument('--dimensions', default='Dimensions.csv', help='Path to Dimensions CSV')
    parser.add_argument('--events', default='Events.csv', help='Path to Events CSV')
    parser.add_argument('--packages', default='Implementation Packages.csv', help='Path to Implementation Packages CSV')
    parser.add_argument('--products', default='Products.csv', help='Path to Products CSV')
    parser.add_argument('--sdk', default='SDKs.csv', help='Path to SDK mapping CSV')
    parser.add_argument('--output-dir', default='output', help='Directory to save output files')
    args = parser.parse_args()

    dims_df = read_csv(args.dimensions)
    events_df = read_csv(args.events)
    packages_df = read_csv(args.packages)
    products_df = read_csv(args.products)
    sdk_df = read_csv(args.sdk)

    # Ensure SDK mapping CSV has a merge key for platform
    if 'platform' not in sdk_df.columns:
        # If SDK CSV uses 'Name' for platform, rename it
        if 'Name' in sdk_df.columns:
            sdk_df = sdk_df.rename(columns={'Name': 'platform'})
        else:
            print(
                "SDK mapping CSV must include a column matching Products.app_platform "
                "(e.g., 'platform' or 'Name')",
                file=sys.stderr
            )
            sys.exit(1)

    # Select products to generate requirements for
    product_names = sorted(products_df['Name'].dropna().unique())
    selected_products = choose_from_list(product_names, 'Select Products')
    filtered_products_df = products_df[products_df['Name'].isin(selected_products)]

    # Select mode
    modes = ['Implementation Packages', 'Direct Events']
    mode_choice = choose_from_list(modes, 'Select mode (Implementation Packages or Direct Events)')
    if not mode_choice:
        print('No mode selected, exiting.', file=sys.stderr)
        sys.exit(1)
    mode = mode_choice[0]

    selected_packages_df = None
    if mode == 'Implementation Packages':
        pkg_opts = sorted(packages_df['Name'].dropna().unique())
        sel_pkgs = choose_from_list(pkg_opts, 'Select Implementation Packages')
        selected_packages_df = packages_df[packages_df['Name'].isin(sel_pkgs)]
        # Extract event IDs from packages
        event_ids = set()
        for _, row in selected_packages_df.iterrows():
            for col in ['Required Events', 'Required Heartbeat Events']:
                s = row.get(col, '')
                if pd.isna(s) or not str(s).strip():
                    continue
                for item in str(s).split(','):
                    m = re.search(r'event(\d+)', item)
                    if m:
                        try:
                            event_ids.add(int(m.group(1)))
                        except ValueError:
                            pass
        selected_events_df = events_df[events_df['Event #'].isin(event_ids)]
    else:
        evt_opts = sorted(events_df['Event Name'].dropna().unique())
        sel_evts = choose_from_list(evt_opts, 'Select Direct Events')
        selected_events_df = events_df[events_df['Event Name'].isin(sel_evts)]

    # Determine required dimensions for selected events
    dim_names = set()
    for _, row in selected_events_df.iterrows():
        s = row.get('Full Integration', '')
        if pd.isna(s) or not str(s).strip():
            s = row.get('MVP Evars And Props', '')
        if pd.isna(s) or not str(s).strip():
            continue
        for item in str(s).replace('\n', ' ').split(','):
            name = item.strip()
            if name:
                dim_names.add(name)
    selected_dims_df = dims_df[dims_df['Name'].isin(dim_names)]

    # Merge products with SDK info
    products_and_sdks_df = filtered_products_df.merge(sdk_df, left_on='app_platform', right_on='platform', how='left')

    # Create output directory
    os.makedirs(args.output_dir, exist_ok=True)

    # Build mapping of events to packages
    event_to_pkgs = {}
    if selected_packages_df is not None:
        for pkg_name, row in selected_packages_df.set_index('Name').iterrows():
            for col in ['Required Events', 'Required Heartbeat Events']:
                s = row.get(col, '')
                if pd.isna(s) or not str(s).strip():
                    continue
                for item in str(s).split(','):
                    m = re.search(r'event(\d+)', item)
                    if m:
                        eid = int(m.group(1))
                        event_to_pkgs.setdefault(eid, []).append(pkg_name)
    # Build mapping of event to its dimensions (full detail)
    dims_by_event = {}
    for _, row in selected_events_df.iterrows():
        eid = row.get('Event #')
        # Collect dimension names from relevant event fields
        dim_names = set()
        for col in ['Full Integration', 'MVP Evars And Props', 'Minimum Viable Evars', 'Minimum Viable Props']:
            if col in row and pd.notna(row[col]):
                # split by comma, strip whitespace
                for item in str(row[col]).replace('\n', ' ').split(','):
                    name = item.strip()
                    if name:
                        dim_names.add(name)
        # Filter dimensions by exact match
        matched = dims_df[dims_df['Name'].isin(dim_names)]
        # If no exact matches, fallback to substring matches
        if matched.empty and dim_names:
            mask = dims_df['Name'].fillna('').apply(
                lambda x: any(name in x for name in dim_names)
            )
            matched = dims_df[mask]
        dims_by_event[eid] = matched
    # Create output directory if needed
    os.makedirs(args.output_dir, exist_ok=True)
    # Generate per-product requirement files
    for _, prod in products_and_sdks_df.iterrows():
        product_name = prod['Name']
        safe_name = sanitize_sheet_name(product_name)
        product_dir = os.path.join(args.output_dir, safe_name)
        os.makedirs(product_dir, exist_ok=True)
        # Filter packages/events per product
        if mode == 'Implementation Packages':
            pkgs_prod = selected_packages_df[
                selected_packages_df.get('Used by product', '').str.contains(product_name, na=False)
            ]
            event_ids_prod = {
                eid for eid, pkgs in event_to_pkgs.items()
                if any(p in pkgs for p in pkgs_prod['Name'].values)
            }
            evts_prod = selected_events_df[selected_events_df['Event #'].isin(event_ids_prod)]
        else:
            pkgs_prod = None
            evts_prod = selected_events_df
        # Generate outputs for this product
        generate_excel_for_product(prod, pkgs_prod, evts_prod, dims_by_event, product_dir, event_to_pkgs)
        prod_df = prod.to_frame().T
        generate_word_for_product(prod_df, pkgs_prod, evts_prod, dims_by_event, product_dir)
        generate_pdf_for_product(prod_df, pkgs_prod, evts_prod, dims_by_event, product_dir)

if __name__ == '__main__':
    main()